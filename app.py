from routes.paddle_webhook import paddle_bp


from flask import Flask, jsonify, request, send_file, g
from flask_cors import CORS
from langchain_openai.chat_models.base import BaseChatOpenAI
from langchain_openai import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_groq import ChatGroq
import re
import json
import uuid
import os
from datetime import datetime
from math import ceil

import tiktoken
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt

from dotenv import load_dotenv
load_dotenv()

from db import get_user_from_token, supabase

app = Flask(__name__)
CORS(app)
app.register_blueprint(paddle_bp)


encoding = tiktoken.encoding_for_model("gpt-3.5-turbo-0125") 


# 设置最大生成次数和字数
MAX_GENERATIONS_PER_DAY = {"free": 10, "pro": 100}
MAX_WORDS_PER_GEN = 500

# CREDIT 单价
CREDIT_VALUE_USD = 0.001

# 设置不同模型的token价格
MODEL_PRICE = {
    "GPT-3.5": { # GPT-3.5-Turbo-0125
        "input_price_per_1m_tokens": 0.5,
        "output_price_per_1m_tokens": 1.5,
        "safety_margin": 0.6
    },
    "GPT-4o-mini": {
        "input_price_per_1m_tokens": 0.15,
        "output_price_per_1m_tokens": 0.6,
        "safety_margin": 1
    },
    "GPT-4o": {
        "input_price_per_1m_tokens": 5,
        "output_price_per_1m_tokens": 15,
        "safety_margin": 0.4
    },
    "GEMINI": {
        "input_price_per_1m_tokens": 0.5,
        "output_price_per_1m_tokens": 1.5,
        "safety_margin": 0.6
    },
    "DEEPSEEK": {
        "input_price_per_1m_tokens": 0.5,
        "output_price_per_1m_tokens": 1.5,
        "safety_margin": 0.5
    },
    "LLAMA3": {
        "input_price_per_1m_tokens": 0.5,
        "output_price_per_1m_tokens": 1.5,
        "safety_margin": 0.6
    },
}


def get_model(model_name):
    if model_name == "GPT-3.5": 
        llm = ChatOpenAI(model_name="gpt-3.5-turbo-0125", max_tokens=1000, temperature=0.2)
    
    if model_name == "GPT-4o-mini":
        llm = ChatOpenAI(model_name="gpt-4o-mini", max_tokens=1000, temperature=0.2)

    elif model_name == "GEMINI":
        llm = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.2)

    elif model_name == "LLAMA3":
        llm = ChatGroq(model="llama3-8b-8192", temperature=0.2)

    elif model_name == "DEEPSEEK": # the deepseek paid API
        llm = BaseChatOpenAI(
            model='deepseek-chat', 
            openai_api_base='https://api.deepseek.com',
            max_tokens=1000,
            api_key=os.getenv("DEEPSEEK_API_KEY"),
            temperature=0.2
        )
    return llm
   

def extract_context_data(context_text, window=1200):

    # 查找选中的文本并获取位置
    #print("context_text:", context_text)
    selected_match = re.search(r"\[\[SELECTED\]\](.*?)\[\[/SELECTED\]\]", context_text, re.DOTALL)
    #print("selected_match:", selected_match)
    cursor_match = re.search(r"\[\[CURSOR\]\]", context_text)
    #print("cursor_match:", cursor_match)

    selected_text = ""

    if selected_match:
        selected_text = selected_match.group(1) 
        selected_start = selected_match.start(0)
        selected_end = selected_match.end(0)

    elif cursor_match:
        selected_start = cursor_match.start()
        selected_end = cursor_match.end()

    else:
        selected_start = None
        selected_end = None

    before_text = context_text[:selected_start][-window:] if selected_start else ""
    after_text = context_text[selected_end:][:window] if selected_end else ""
    #print("before_text:", before_text)
    #print("after_text:", after_text)
    return selected_text, before_text, after_text


def construct_prompt(mode, context_text, tone, style, audience, custom_prompt):

    selected_text, before_text, after_text = extract_context_data(context_text)
    
    prompt = "You are a professional writing assistant.\n"
    
    if mode == "polish":
        prompt += "Your task is to improve the clarity, style, and correctness of the SELECTION.\n"
        prompt += "Consider the surrounding CONTEXT for better coherence.\n"

    elif mode == "expand":
        prompt += "Your task is to expand the SELECTION by adding more details, depth, and elaboration.\n"
        prompt += "Ensure the expanded content integrates smoothly into the CONTEXT.\n"

    elif mode == "continue":
        prompt += "Your task is to continue writing naturally from the SELECTION\n"
        prompt += "Ensure the continuation explores new ideas while maintaining consistency with the CONTEXT.\n"
        prompt += "Avoid repeating information already covered in the CONTEXT.\n"

    prompt += f"\n# CONTEXT and SELECTION:\n"
    
    if before_text:
        prompt += f"\n## Context Before the Selection:\n{before_text}\n"
    
    if selected_text:
        prompt += f"\n## SELECTION:\n{selected_text}\n"
    else:
        prompt += f"\n## SELECTION\n"

    if mode == "continue":
        prompt += f"(continue writing here...)\n"
    elif mode == "expand":
        prompt += f"(expand the SELECTION text with more details...)\n"
    elif mode == "polish":
        prompt += f"(polish the SELECTION text...)\n"
    
    if after_text:
        prompt += f"\n## Context After the Selection:\n{after_text}\n"

    prompt += f"\nEnsure the generated text flows naturally and seamlessly into the surrounding context. Default to align with the context language."

    if any([tone, style, audience, custom_prompt]):
        prompt += "\nPlease follow the instructions below to meet user's specific needs."

    if tone:
        prompt += f"\nTone: {tone}"
    
    if style:
        prompt += f"\nStyle: {style}"
    
    if audience:
        prompt += f"\nTarget audience: {audience}"
    
    if custom_prompt:
        prompt += f"\n{custom_prompt}"

    prompt += f"\nReturn only the continuation text without any introductory or concluding sentence."

    return prompt


def calculate_credit_spent(model_name, input_tokens, output_tokens):
    if model_name not in MODEL_PRICE:
        raise ValueError(f"Model {model_name} unknown.")
    pricing = MODEL_PRICE.get(model_name)

    input_cost = input_tokens * pricing.get("input_price_per_1m_tokens") / 1000000
    output_cost = output_tokens * pricing.get("output_price_per_1m_tokens") / 1000000
    total_cost = input_cost + output_cost
    total_cost_with_safety_margin = total_cost * (1 + pricing.get("safety_margin")) # add a safety margin for profit

    credits = ceil(total_cost_with_safety_margin / CREDIT_VALUE_USD * 10) / 10 # round up to 1 decimal place
    return credits




@app.before_request # register a request handler
def before_request():
    if request.endpoint in ['public_endpoint', 'login_endpoint', 'ping', 'paddle.handle_paddle_webhook']: # skip authentication for public endpoints
        return
    
    user_id = get_user_from_token()

    if user_id:
        g.user_id = user_id
    else:
        g.user_id = None

@app.route('/ping', methods=['GET'])
def ping():
    return jsonify({"message": "pong"})

@app.route("/get_tokens", methods=["GET"])
def get_tokens():
    if not g.user_id:
        return jsonify({"error": "User not authenticated"}), 401
    user_id = g.user_id

    response = supabase.table("users").select("tokens").eq("auth_id", user_id).execute()
    if not response.data:
        return jsonify({"error": "User not found"}), 40
    current_tokens = response.data[0]["tokens"]      

    return jsonify({"tokens": current_tokens})

@app.route("/feedback", methods=["POST"])
def feedback():
    if not g.user_id:
        return jsonify({"error": "User not authenticated"}), 401
    user_id = g.user_id

    data = request.json
    message = data.get("message").strip()
    email = data.get("email").strip()

    if message: # insert feedback to database
        response = supabase.table("feedback").insert({
            "auth_id": user_id,
            "message": message,
            "email": email,
        }).execute()

        if not response.data:
            return jsonify({"error": "Failed to insert feedback"}), 500

    return jsonify({"message": "success"})

@app.route("/billing", methods=["GET"])
def get_billing_info():
    
    if not g.user_id:
        return jsonify({"error": "User not authenticated"}), 401
    user_id = g.user_id

    # get current tokens from user table
    response = supabase.table("users").select("tokens").eq("auth_id", user_id).execute()
    if not response.data:
        return jsonify({"error": "User not found"}), 404
    current_tokens = response.data[0]["tokens"]  


    # get history purchases
    response = supabase.table("transactions").select("*").eq("auth_id", user_id).execute()
    if not response.data:
        transactions = []
    else:
        transactions = response.data


    return jsonify({
        "balance": current_tokens,
        "transactions": [
            {
                "amount": t.get("amount"),
                "token_amount": t.get("token_amount"),
                "tx_type": t.get("type"), # top-up, subscription, etc.
                "price_tag": t.get("price_tag"), # the selected sub price product name, or plan name 
                "quantity": t.get("quantity"), # how many of the price product or plan is purchased
                "status": t.get("status"), 
                "tx_id": t.get("paddle_transaction_id"),
                "created_at": datetime.strptime(t.get("created_at"), "%Y-%m-%dT%H:%M:%S.%f%z").strftime("%Y-%m-%d %H:%M:%S")
            }
            for t in transactions
        ]
    })


@app.route('/generate', methods=['POST'])
def generate():
    if not g.user_id:
        return jsonify({"error": "User not authenticated"}), 401
    user_id = g.user_id

    data = request.json
    generate_mode = data.get("selected_mode", "continue").strip()
    context_text = data.get("context_text").strip()

    tone = data.get("tone").strip()
    style = data.get("style").strip()
    audience = data.get("audience").strip()
    custom_prompt = data.get("customer_prompt", "").strip()

    model_name = data.get("model")

    # get user info (type, subscription tier, and tokens)
    response = supabase.table("users").select("role", "subscription_tier", "tokens").eq("auth_id", user_id).execute()
    user_data = response.data
    if not user_data:
        return jsonify({"error": "User not found"}), 404
    
    user = user_data[0]
    #role = user.get("role")
    subscription_tier = user.get("subscription_tier")
    credits = user.get("tokens") # 代币数量
    daily_count = user.get("daily_gen_count")
    last_gen_date = user.get("last_gen_date")
    #print("the current last gen date is:", last_gen_date, type(last_gen_date))
    
    # limit the API call based on rate limits, differs based on the subscription tier
    today_date = datetime.today().date().isoformat()
    if last_gen_date != today_date:
        daily_count = 0
    elif daily_count >= MAX_GENERATIONS_PER_DAY.get(subscription_tier, 10):    
        return jsonify({"error": "Generation limit exceeded for today"}), 402
    
    # prepare prompt
    final_prompt = construct_prompt(generate_mode, context_text, tone, style, audience, custom_prompt)
    print("final_prompt:", final_prompt)

    # limit the API call based on token limits
    tokens_prompt = len(encoding.encode(final_prompt))  # a rough estimate of input tokens
    credits_needed = calculate_credit_spent(model_name, input_tokens=tokens_prompt, output_tokens=MAX_WORDS_PER_GEN) # a roungh estimate of credits needed
    if credits_needed > credits:
        return jsonify({"error": "Token limit exceeded"}), 402
    
    llm = get_model(model_name)

    try:
        response = llm.invoke(final_prompt)        
        print("response:", response)
        text = response.content
        
        input_tokens = response.usage_metadata['input_tokens']
        output_tokens = response.usage_metadata['output_tokens']

        # TESTING on mockup response
        #text = """this is mockup LLM response for testing purposes.\n\nThe generated **bold text** , and *italic text*, as well as the ~~strikethrough text~~ , lastly the [link text](https://www.google.com) will be here. \nAlso have some markdown formatting for testing as well.\n\n- list item 1\n- list item 2\n- list item 3\n1. ordered list item 1\n2. ordered list item 2\n"""
    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500
      
    # update user tokens and daily count to supabase table
    credits_spent = calculate_credit_spent(model_name, input_tokens, output_tokens)

    new_credits = credits - credits_spent # float
    
    response = supabase.table("users").update({
        "tokens": new_credits,   # update in db float type
        "daily_gen_count": int(daily_count + 1),
        "last_gen_date": today_date
    }).eq("auth_id", user_id).execute()
    
    # TODO: update usage_daily 表  // 每日局和写入

    # check response success or not
    if not response.data:
        return jsonify({"error": response.error['message']}), 502

    return jsonify({"generated_text": text, "tokens": int(new_credits)}) # returned credits in integer type

@app.route("/ensure_user", methods=["POST"])
def ensure_user():
    if not g.user_id:
        return jsonify({"error": "User not authenticated"}), 401
    user_id = g.user_id

    # 查询 users 表
    existing_user = supabase.table("users").select("*").eq("auth_id", user_id).execute()

    if not existing_user.data:  # 如果 users 表里没有该用户
        tokens = 1000 # 新用户默认获赠代币
        role = "user"
        subscription_tier = "free"  
        supabase.table("users").insert([
            {"auth_id": user_id, "role": role, "subscription_tier": subscription_tier, "tokens": tokens}
        ]).execute()
    else:
        tokens = existing_user.data[0]["tokens"]
        role = existing_user.data[0]["role"]
        subscription_tier = existing_user.data[0]["subscription_tier"]

    return jsonify({"tokens": tokens, "role": role, "subscription_tier": subscription_tier})


#document_storage = {} # temporary storage for document content
@app.route('/save_document', methods=['POST'])
def save_document():
    data = request.get_json()
    content = data.get("content")
    #print(f"content: {content}")
    #print(f"content type: {type(content)}")

    if not g.user_id:
        return jsonify({"error": "User not authenticated"}), 401
    user_id = g.user_id
    #print(f"Saving document for user {user_id}")

    response = supabase.table("documents").upsert([
        {"user_id": user_id, "content": content}
    ], on_conflict=["user_id"]).execute()
    #document_storage[user_id] = content

    if response.data:
        #print(f"Document [{content}] saved for user {user_id}")
        return jsonify({"message": "Content saved successfully"})
    
    return jsonify({"error": response.error['message']}), 500


@app.route('/get_document', methods=['GET'])
def get_document():
    
    if not g.user_id:
        return jsonify({"error": "User not authenticated"}), 401
    user_id = g.user_id

    response = supabase.table("documents").select("content").eq("user_id", user_id).execute()

    if not response.data:
        return jsonify({"error": response.error['message']}), 500
    
    content = response.data[0]['content']


    #content = document_storage.get(user_id, "")
    #print(f"Retrieved document {content} for user {user_id}")

    return jsonify({"content": content})


@app.route('/export_document', methods=['GET'])
def export_document():
    if not g.user_id:
        return jsonify({"error": "User not authenticated"}), 401
    user_id = g.user_id
    response = supabase.table("documents").select("content").eq("user_id", user_id).execute()

    if not response.data:
        return jsonify({"error": response.error['message']}), 500
    
    content = response.data[0]['content']

    print(f"Exporting document for user {user_id}")

    try:
        lexical_content = json.loads(content)
    except Exception as e:
        return jsonify({"error": "Invalid JSON content"}), 400
    

    # ✅ 生成 Word 文档
    word_file_path = f"documents/{user_id}.docx"
    document = generate_word_from_lexical_content(lexical_content)
    document.save(word_file_path)

    return send_file(word_file_path, as_attachment=True, download_name=f"user_{user_id}.docx")



# ==== Word 文档生成器 ==== # TODO: move to a separate file

def generate_word_from_lexical_content(lexical_content):
    document = Document()

    # 设置全局默认字体
    set_default_font(document)

    # 确保 lexical_content 是字典格式并且包含根节点
    if isinstance(lexical_content, dict) and "root" in lexical_content:
        for node in lexical_content["root"]["children"]:
            if node["type"] == "paragraph":
                para = document.add_paragraph()
                add_text_from_node(para, node)
            elif node["type"] == "heading":
                # 标题处理（例如 h1 转换为 Heading 1）
                para = document.add_paragraph(style='Heading 1' if node.get("tag") == "h1" else 'Heading 2')
                add_text_from_node(para, node)
            elif node["type"] == "unordered-list":
                # 无序列表处理
                add_list_from_node(document, node, unordered=True)
            elif node["type"] == "ordered-list":
                # 有序列表处理
                add_list_from_node(document, node, unordered=False)
    else:
        raise ValueError("Invalid lexical content format")

    return document

def set_default_font(document):
    # 设置全局字体样式（中文宋体，英文 Calibri）
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    
    # 对中文进行单独设置
    font_element = style.element
    r_pr = font_element.get_or_add_rPr()
    
    # 设置中文字体为宋体
    chinese_font = OxmlElement('w:rFonts')
    chinese_font.set(qn('w:ascii'), 'Calibri')
    chinese_font.set(qn('w:eastAsia'), 'SimSun')  # 中文设置为宋体
    r_pr.append(chinese_font)

def add_text_from_node(paragraph, node):
    # 对每个段落或文本节点处理文本内容和格式
    for child in node["children"]:
        text = child["text"]
        run = paragraph.add_run(text)
        
        # 应用文本样式
        if child.get("format") == 1:  # 粗体
            run.bold = True
        if child.get("format") == 2:  # 斜体
            run.italic = True
        if child.get("format") == 4:  # 下划线
            run.underline = True
        if child.get("format") == 5:  # 删除线
            run.strike = True

def add_list_from_node(document, node, unordered=True):
    # 添加列表
    for child in node["children"]:
        if unordered:
            para = document.add_paragraph(style='List Bullet')
        else:
            para = document.add_paragraph(style='List Number')
        add_text_from_node(para, child)



if __name__ == '__main__':
    app.run(debug=True)
